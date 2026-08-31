"""Build a concise Word summary for the two xingrong_3 forecast models."""

import csv
import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[3]
OUT_DIR = ROOT / "reports"
ASSET_DIR = ROOT / "runtime" / "reports" / "model_training_summary_20260818_assets"
OUT = OUT_DIR / "兴蓉三号炉_预测模型训练成果总结_20260818.docx"
PYTHON = ""

FONT = r"C:\Windows\Fonts\msyh.ttc"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MUTED = "666666"
TABLE_WIDTH_DXA = 9360


def chinese_font(size, bold=False):
    try:
        return ImageFont.truetype(FONT, size, index=1 if bold else 0)
    except OSError:
        return ImageFont.load_default()


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    for grid_col, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_para(paragraph, before=0, after=6, line=1.10, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if keep:
        fmt.keep_with_next = True


def add_text(paragraph, text, size=11, color=INK, bold=False, italic=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return run


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    size, color, before, after = {1: (16, BLUE, 16, 8), 2: (13, BLUE, 12, 6)}[level]
    set_para(p, before=before, after=after, line=1.10, keep=True)
    add_text(p, text, size=size, color=color, bold=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, before=2, after=8, line=1.0)
    add_text(p, text, size=9, color=MUTED, italic=True)


def make_trend_chart(rows, out_path):
    width, height = 1600, 850
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = chinese_font(38, True)
    label_font = chinese_font(25)
    small_font = chinese_font(22)
    draw.text((72, 42), "关键目标的 10 分钟趋势方向一致率", fill="#0B2545", font=title_font)
    draw.text((72, 96), "仅统计真实变化较明显的窗口；数值越高表示预测涨跌方向越可靠", fill="#666666", font=small_font)
    left, right, top, bottom = 420, 1450, 170, 760
    for pct in range(0, 101, 20):
        x = left + (right-left) * pct / 100
        draw.line((x, top, x, bottom), fill="#E5E7EB", width=2)
        draw.text((x-16, bottom+14), str(pct), fill="#666666", font=small_font)
    draw.text((right-28, bottom+52), "%", fill="#666666", font=small_font)
    ordered = sorted(rows, key=lambda r: float(r["trend_direction_accuracy"]))
    spacing = (bottom-top) / len(ordered)
    for i, row in enumerate(ordered):
        y = top + int(i * spacing + 14)
        value = float(row["trend_direction_accuracy"]) * 100
        name = row["name"].replace("SO₂", "SO2")
        model = row["model"]
        color = "#0077BB" if model == "子模型1" else "#EE7733"
        draw.rounded_rectangle((left, y, left + (right-left)*value/100, y+31), radius=8, fill=color)
        draw.text((70, y-3), f"{name}  ({model})", fill="#1F2937", font=label_font)
        draw.text((left + (right-left)*value/100 + 15, y-2), f"{value:.1f}%", fill="#1F2937", font=label_font)
    draw.rectangle((72, 798, 95, 821), fill="#0077BB")
    draw.text((105, 792), "子模型1：炉膛与燃烧", fill="#4B5563", font=small_font)
    draw.rectangle((365, 798, 388, 821), fill="#EE7733")
    draw.text((398, 792), "子模型2：烟气净化与排放", fill="#4B5563", font=small_font)
    image.save(out_path)


def load_json(path):
    return json.loads(path.read_text(encoding="utf-8"))


def load_rows(path):
    with path.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def add_table(doc, headers, data, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, label in zip(header.cells, headers):
        shade(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para(p, before=0, after=0, line=1.0)
        add_text(p, label, size=font_size, color=DARK_BLUE, bold=True)
    for record in data:
        cells = table.add_row().cells
        for cell, value in zip(cells, record):
            p = cell.paragraphs[0]
            set_para(p, before=0, after=0, line=1.0)
            add_text(p, str(value), size=font_size, color="222222")
    return table


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para(footer, before=0, after=0, line=1.0)
    add_text(footer, "兴蓉三号炉 · 预测模型训练成果总结 · 2026-08-18", size=8.5, color=MUTED)
    return doc


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    m1 = load_json(ROOT / "offline/artifacts/forecast/training_ready/model1_bundle_source/metadata.json")
    m2 = load_json(ROOT / "offline/artifacts/forecast/training_ready/model2_bundle_source/metadata.json")
    metrics = load_rows(ROOT / "runtime/runs/model_forecast_review/20260817/key_target_metrics_sampled.csv")
    chart_path = ASSET_DIR / "trend_direction_accuracy.png"
    make_trend_chart(metrics, chart_path)
    shutil.copy2(ROOT / "offline/artifacts/forecast/training_ready/model1_bundle_source/loss_history.png", ASSET_DIR / "model1_loss_history.png")
    shutil.copy2(ROOT / "offline/artifacts/forecast/training_ready/model2_bundle_source/loss_history.png", ASSET_DIR / "model2_loss_history.png")

    doc = setup_document()
    p = doc.add_paragraph()
    set_para(p, before=0, after=3, line=1.0)
    add_text(p, "离线训练结果摘要", size=10, color=BLUE, bold=True)
    p = doc.add_paragraph()
    set_para(p, before=0, after=7, line=1.0)
    add_text(p, "兴蓉三号炉预测模型训练成果总结", size=22, color=INK, bold=True)
    p = doc.add_paragraph()
    set_para(p, before=0, after=14, line=1.10)
    add_text(p, "两子模型：炉膛与燃烧过程、烟气净化与烟囱排放", size=12, color=MUTED)
    add_text(p, "  |  版本：pred600-v20260809-r1  |  2026-08-18", size=10, color=MUTED)

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [TABLE_WIDTH_DXA])
    shade(callout.cell(0, 0), "F4F6F9")
    p = callout.cell(0, 0).paragraphs[0]
    set_para(p, before=0, after=0, line=1.15)
    add_text(p, "核心结论：", size=10.5, color=DARK_BLUE, bold=True)
    add_text(p, "两个模型已完成离线训练与结构校验，可提供未来 10 分钟的工况基线。烟气指标的趋势识别总体较好；炉膛侧部分变量（尤其主蒸汽流量）更适合作为辅助预警，不宜单独触发控制动作。", size=10.5, color=INK)

    add_heading(doc, "1. 训练数据与任务定义")
    data_rows = [
        ["子模型1", "炉膛与燃烧过程", "2025-05-11 至 2026-05-14", "300", "5", "186,708 / 38,120 / 36,591"],
        ["子模型2", "烟气净化与烟囱排放", "2025-05-11 至 2026-05-14", "109", "5", "9,598,582 / 1,992,398 / 1,936,787"],
    ]
    add_table(doc, ["模型", "覆盖环节", "数据时间范围", "输入变量", "预测目标", "训练/验证/测试窗口"], data_rows,
              [930, 1760, 1970, 880, 820, 3000], 8.8)
    p = doc.add_paragraph()
    set_para(p, before=4, after=6, line=1.10)
    add_text(p, "统一任务口径：", size=10, color=DARK_BLUE, bold=True)
    add_text(p, "1 秒粒度；使用过去 20 分钟（1,200 点）预测未来 10 分钟（600 点）。两套数据按时间先后划分训练、验证和测试集，避免使用未来数据预测过去。", size=10, color=INK)

    add_heading(doc, "2. 模型与训练参数")
    param_rows = [
        ["网络", "Crossformer", "Crossformer"],
        ["隐藏维度 / 注意力头", "192 / 6", "128 / 4"],
        ["编码层 / 解码层", "2 / 1", "2 / 1"],
        ["训练轮次", "8", "8"],
        ["批量大小", "128", "256"],
        ["学习率", "2e-05", "3e-05"],
        ["训练设备", "8 GPU 并行", "8 GPU 并行"],
    ]
    add_table(doc, ["参数", "子模型1", "子模型2"], param_rows, [2250, 3555, 3555], 9.2)
    p = doc.add_paragraph()
    set_para(p, before=5, after=2, line=1.10)
    add_text(p, "目标变量：", size=10, color=DARK_BLUE, bold=True)
    add_text(p, "子模型1预测主蒸汽流量、省煤器出口氧量、氧量、中上部炉膛温度和垃圾层厚度；子模型2预测 CO、HCl、NOx、SO₂ 和烟尘。", size=10, color=INK)

    add_heading(doc, "3. 训练收敛情况")
    fig_table = doc.add_table(rows=1, cols=2)
    set_table_geometry(fig_table, [4680, 4680])
    for cell, path in zip(fig_table.rows[0].cells, [ASSET_DIR / "model1_loss_history.png", ASSET_DIR / "model2_loss_history.png"]):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].add_run().add_picture(str(path), width=Inches(3.05))
    add_caption(doc, "图 1  两个子模型在 8 个训练轮次中的训练、验证和测试损失变化（越低越好）")
    p = doc.add_paragraph()
    set_para(p, before=0, after=6, line=1.15)
    add_text(p, "解读：", size=10, color=DARK_BLUE, bold=True)
    add_text(p, "两个模型的损失均在前几轮快速下降，后期趋于稳定；未出现持续恶化的验证损失。子模型1的最终测试损失为 0.269，子模型2为 0.0086（两者分别在自身归一化训练口径下计算，不能直接跨模型比较数值大小）。", size=10, color=INK)

    add_heading(doc, "4. 预测结果（工程单位）")
    p = doc.add_paragraph()
    set_para(p, before=0, after=5, line=1.10)
    add_text(p, "评估方法：", size=10, color=DARK_BLUE, bold=True)
    add_text(p, "从已保存测试输出中，对每个模型均匀抽取 10,000 个窗口；将预测值按各模型 scaler 还原至工程单位。MAE 表示平均绝对误差，数值越小越好。", size=10, color=INK)
    result_rows = []
    for row in metrics:
        result_rows.append([
            row["model"], row["name"], row["unit"],
            f"{float(row['mae_all_10min']):.3f}", f"{float(row['rmse_all_10min']):.3f}",
            f"{float(row['trend_direction_accuracy'])*100:.1f}%",
        ])
    add_table(doc, ["模型", "目标", "单位", "10 分钟 MAE", "10 分钟 RMSE", "趋势方向一致率"], result_rows,
              [980, 2100, 1130, 1670, 1690, 1790], 8.6)

    doc.add_page_break()
    add_heading(doc, "5. 简要分析与适用边界")
    doc.add_picture(str(chart_path), width=Inches(6.45))
    add_caption(doc, "图 2  关键目标的 10 分钟趋势方向一致率（抽样测试）")
    bullets = [
        "烟气侧：CO 和 SO₂ 的趋势方向一致率分别为 84.1% 和 87.9%，适合用于未来排放风险的提前识别；NOx 为 74.9%，可作为辅助依据。",
        "炉膛侧：中上部炉膛温度和垃圾层厚度的趋势识别相对较好；主蒸汽流量为 47.4%，其未来涨跌不应作为单一控制触发条件。",
        "曲线形态：原始 1 秒预测存在高频波动偏多的现象。离线试验表明，按连续 10 秒均值聚合后，误差基本不变，同时更贴近实际控制的 10 秒更新节奏。",
        "应用定位：模型适合作为后续规则、历史策略报告和 MPC 的“未来工况基线”。具体控制动作仍应受人工规则、设备边界、限值与安全裁决共同约束。",
    ]
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        set_para(p, before=0, after=4, line=1.12)
        for run in p.runs:
            set_run_font(run, size=10, color=INK)
        add_text(p, item, size=10, color=INK)

    add_heading(doc, "6. 结论")
    p = doc.add_paragraph()
    set_para(p, before=0, after=5, line=1.15)
    add_text(p, "本轮训练完成了面向兴蓉三号炉的两子模型建设，模型包已通过输入输出结构核验。", size=10.5, color=INK)
    add_text(p, "当前结果支持离线回放、预测展示和控制辅助研究；尚不代表现场闭环效果已验证。", size=10.5, color=DARK_BLUE, bold=True)
    p = doc.add_paragraph()
    set_para(p, before=0, after=0, line=1.1)
    add_text(p, "建议后续以 10 秒聚合后的预测轨迹进入诊断与控制决策层，并在现场数据接入后继续验证越限预警、趋势判断和控制收益。", size=10.5, color=INK)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
